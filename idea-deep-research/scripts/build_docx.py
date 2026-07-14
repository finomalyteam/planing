#!/usr/bin/env python3
"""แปลงรายงาน Markdown เป็น Word (.docx)

การใช้งาน: python build_docx.py <input.md> <output.docx>
ต้องมี python-docx: pip install python-docx

รองรับ: heading (# ถึง ####), bullet (- หรือ *), numbered list (1.),
ตาราง pipe (| a | b |), **ตัวหนา**, และ blockquote (>)
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ไม่พบ python-docx — ติดตั้งก่อนด้วย: pip install python-docx")
    sys.exit(1)


def add_runs(paragraph, text: str) -> None:
    """แตก **bold** เป็น run แยก"""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1


def is_table_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", line.strip())) and "-" in line


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 1

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    if not src.is_file():
        print(f"ไม่พบไฟล์: {src}")
        return 1

    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # ตาราง: เก็บบรรทัด | ติดกันทั้งก้อน
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_sep(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        para = table.cell(r, c).paragraphs[0]
                        add_runs(para, cell)
                        if r == 0:
                            for run in para.runs:
                                run.bold = True
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            doc.add_heading(m.group(2), level=len(m.group(1)))
        elif re.match(r"^[-*]\s+", stripped):
            add_runs(
                doc.add_paragraph(style="List Bullet"),
                re.sub(r"^[-*]\s+", "", stripped),
            )
        elif re.match(r"^\d+\.\s+", stripped):
            add_runs(
                doc.add_paragraph(style="List Number"),
                re.sub(r"^\d+\.\s+", "", stripped),
            )
        elif stripped.startswith(">"):
            add_runs(
                doc.add_paragraph(style="Intense Quote"),
                stripped.lstrip("> "),
            )
        else:
            add_runs(doc.add_paragraph(), stripped)
        i += 1

    doc.save(dst)
    print(f"สร้างแล้ว: {dst}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
